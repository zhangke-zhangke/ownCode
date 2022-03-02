# coding=utf-8
import re, os
import jieba.posseg as pseg
import jieba
import time
from py2neo import Graph, Node, Relationship,NodeMatcher
from docx import Document

class ExtractEvent:
    def __init__(self):
        self.map_dict = self.load_mapdict()
        self.minlen = 2
        self.maxlen = 30
        self.keywords_num = 20
        self.limit_score = 10
        self.IP = "(([NERMQ]*P*[ABDP]*)*([ABDV]{1,})*([NERMQ]*)*([VDAB]$)?([NERMQ]*)*([VDAB]$)?)*"
        self.IP = "([NER]*([PMBQADP]*[NER]*)*([VPDA]{1,}[NEBRVMQDA]*)*)"
        self.MQ = '[DP]*M{1,}[Q]*([VN]$)?'
        self.VNP = 'V*N{1,}'
        self.NP = '[NER]{1,}'
        self.REN = 'R{2,}'
        self.VP = 'P?(V|A$|D$){1,}'
        self.PP = 'P?[NERMQ]{1,}'
        self.SPO_n = "n{1,}"
        self.SPO_v = "v{1,}"
        self.stop_tags = {'u', 'wp', 'o', 'y', 'w', 'f', 'u', 'c', 'uj', 'nd', 't', 'x'}
        self.combine_words = {"首先", "然后", "之前", "之后", "其次", "接着"}

    """构建映射字典"""

    def load_mapdict(self):
        tag_dict = {
            'B': 'b'.split(),  # 时间词
            'A': 'a d'.split(),  # 时间词
            'D': "d".split(),  # 限定词
            'N': "n j s zg en l r".split(),  # 名词
            "E": "nt nz ns an ng".split(),  # 实体词
            "R": "nr".split(),  # 人物
            'G': "g".split(),  # 语素
            'V': "vd v va i vg vn g".split(),  # 动词
            'P': "p f".split(),  # 介词
            "M": "m t".split(),  # 数词
            "Q": "q".split(),  # 量词
            "v": "V".split(),  # 动词短语
            "n": "N".split(),  # 名词介宾短语
        }
        map_dict = {}
        for flag, tags in tag_dict.items():
            for tag in tags:
                map_dict[tag] = flag
        return map_dict

    """根据定义的标签,对词性进行标签化"""

    def transfer_tags(self, postags):
        # tags = [self.map_dict.get(tag[:2], 'W') for tag in postags]
        # 在自定义字典里查找postags里每个tag对应的标签，如果没查到则认为是W
        tags = [self.map_dict.get(tag, 'W') for tag in postags]

        return ''.join(tags)

    """抽取出指定长度的ngram"""

    def extract_ngram(self, pos_seq, regex):
        ss = self.transfer_tags(pos_seq)

        def gen():
            for s in range(len(ss)):
                for n in range(self.minlen, 1 + min(self.maxlen, len(ss) - s)):
                    e = s + n
                    substr = ss[s:e]
                    if re.match(regex + "$", substr):
                        yield (s, e)

        return list(gen())

    '''抽取ngram'''

    def extract_sentgram(self, pos_seq, regex):
        ss = self.transfer_tags(pos_seq)

        def gen():
            for m in re.finditer(regex, ss):
                yield (m.start(), m.end())

        return list(gen())

    """指示代词替换，消解处理"""
    #                          词      词性      人名
    def cite_resolution(self, words, postags, persons):
        #print(f"处理词为:{words},词性为:{postags},人名为:{persons}")
        if not persons and 'r' not in set(postags):
            #print("1")
            return words, postags
        elif persons and 'r' in set(postags):
            #print("2")
            cite_index = postags.index('r')
            if words[cite_index] in {"其", "他", "她", "我"}:
                words[cite_index] = persons[-1]
                postags[cite_index] = 'nr'
        elif 'r' in set(postags):
            #("3")
            cite_index = postags.index('r')
            if words[cite_index] in {"为何", "何", "如何"}:
                postags[cite_index] = 'w'
        # 判断首句既有人名又有代词噪声
        drop_r_index = []
        if 'r' and 'nr' in set(postags):
            #print("4")
            for i in range(len(words)):
                if postags[i] == 'r':
                    drop_r_index.append(i)
        for i in drop_r_index:
            del words[i]
            del postags[i]

        return words, postags

    """抽取量词性短语"""

    def extract_mqs(self, wds, postags):
        phrase_tokspans = self.extract_sentgram(postags, self.MQ)
        if not phrase_tokspans:
            return []
        phrases = [''.join(wds[i[0]:i[1]]) for i in phrase_tokspans]
        return phrases

    '''抽取动词性短语'''

    def get_ips(self, wds, postags):
        ips = []
        phrase_tokspans = self.extract_sentgram(postags, self.IP)
        #print(phrase_tokspans)
        if not phrase_tokspans:
            return []
        phrases = [''.join(wds[i[0]:i[1]]) for i in phrase_tokspans]
        phrase_postags = [''.join(postags[i[0]:i[1]]) for i in phrase_tokspans]
        for phrase, phrase_postag_ in zip(phrases, phrase_postags):
            if not phrase:
                continue
            #                                               数词              量词               形容词           时间
            phrase_postags = ''.join(phrase_postag_).replace('m', '').replace('q', '').replace('a', '').replace('t', '')
            #                       名词                              简称略语        startswith用来判断是否以指定字符串开头
            if phrase_postags.startswith('n') or phrase_postags.startswith('j'):
                has_subj = 1
            else:
                has_subj = 0
            ips.append((has_subj, phrase))
        return ips

    """分短句处理"""

    def split_short_sents(self, text):
        return [i for i in re.split(r'[,，]', text) if len(i) > 2]

    """分段落"""

    def split_paras(self, text):
        return [i for i in re.split(r'[\n\r]', text) if len(i) > 4]

    """分长句处理"""

    def split_long_sents(self, text):
        return [i for i in re.split(r'[;。:； ：？?!！【】▲丨|]', text) if len(i) > 4]

    """移出噪声数据"""

    def remove_punc(self, text):
        text = text.replace('\u3000', '').replace("'", '').replace('“', '').replace('”', '').replace('▲', '').replace(
            '” ', "”")
        tmps = re.findall('[\(|（][^\(（\)）]*[\)|）]', text)
        for tmp in tmps:
            text = text.replace(tmp, '')
        return text

    """保持专有名词"""

    def zhuanming(self, text):
        books = re.findall('[<《][^《》]*[》>]', text)
        return books

    """对人物类词语进行修正"""

    def modify_nr(self, wds, postags):
        phrase_tokspans = self.extract_sentgram(postags, self.REN)
        wds_seq = ' '.join(wds)
        pos_seq = ' '.join(postags)
        if not phrase_tokspans:
            return wds, postags
        else:
            wd_phrases = [' '.join(wds[i[0]:i[1]]) for i in phrase_tokspans]
            postag_phrases = [' '.join(postags[i[0]:i[1]]) for i in phrase_tokspans]
            for wd_phrase in wd_phrases:
                tmp = wd_phrase.replace(' ', '')
                wds_seq = wds_seq.replace(wd_phrase, tmp)
            for postag_phrase in postag_phrases:
                pos_seq = pos_seq.replace(postag_phrase, 'nr')
        words = [i for i in wds_seq.split(' ') if i]
        postags = [i for i in pos_seq.split(' ') if i]
        return words, postags

    """对人物类词语进行修正"""

    def modify_duplicate(self, wds, postags, regex, tag):
        phrase_tokspans = self.extract_sentgram(postags, regex)
        wds_seq = ' '.join(wds)
        pos_seq = ' '.join(postags)
        if not phrase_tokspans:
            return wds, postags
        else:
            wd_phrases = [' '.join(wds[i[0]:i[1]]) for i in phrase_tokspans]
            postag_phrases = [' '.join(postags[i[0]:i[1]]) for i in phrase_tokspans]
            for wd_phrase in wd_phrases:
                tmp = wd_phrase.replace(' ', '')
                wds_seq = wds_seq.replace(wd_phrase, tmp)
            for postag_phrase in postag_phrases:
                pos_seq = pos_seq.replace(postag_phrase, tag)
        words = [i for i in wds_seq.split(' ') if i]
        postags = [i for i in pos_seq.split(' ') if i]
        return words, postags

    '''对句子进行分词处理'''

    def cut_wds(self, sent):
        WordList = ["一星四射队", "中科院微电子所", "小聋瞎队", "打完就解散队", "女生1队", "2021年", "5月29日", "糊人不唬人队", "男子队伍", "女子队伍","动车","金融反洗钱"]
        for UserWord in WordList:
            jieba.add_word(UserWord)
        wds = list(pseg.cut(sent))
        postags = [w.flag for w in wds]
        words = [w.word for w in wds]
        return self.modify_nr(words, postags)

    """移除噪声词语"""

    def clean_wds(self, words, postags):
        wds = []
        poss = []
        # for wd, postag in zip(words, postags):
        #     print(postags[0].lower())
        #     if postag[0].lower() in self.stop_tags:
        #         continue
        #     wds.append(wd)
        #     poss.append(postag[:2])
        # return wds, poss
        for i in range(len(words)):
            if postags[i].lower() in self.stop_tags:
                continue
            else:
                wds.append(words[i])
                poss.append(postags[i])
        return wds, poss

    """检测是否成立, 肯定需要包括名词"""

    def check_flag(self, postags):
        if not {"v", 'a', 'i'}.intersection(postags):
            return 0
        return 1

    """识别出人名实体"""

    def detect_person(self, words, postags):
        persons = []
        for wd, postag in zip(words, postags):
            if postag == 'nr':
                persons.append(wd)
        return persons

    """识别出名词性短语"""

    def get_nps(self, wds, postags):
        phrase_tokspans = self.extract_sentgram(postags, self.NP)
        if not phrase_tokspans:
            return [], []
        phrases_np = [''.join(wds[i[0]:i[1]]) for i in phrase_tokspans]
        return phrase_tokspans, phrases_np

    """识别出介宾短语"""

    def get_pps(self, wds, postags):
        phrase_tokspans = self.extract_sentgram(postags, self.PP)
        if not phrase_tokspans:
            return [], []
        phrases_pp = [''.join(wds[i[0]:i[1]]) for i in phrase_tokspans]
        return phrase_tokspans, phrases_pp

    """识别出动词短语"""

    def get_vps(self, wds, postags):
        phrase_tokspans = self.extract_sentgram(postags, self.VP)
        if not phrase_tokspans:
            return [], []
        phrases_vp = [''.join(wds[i[0]:i[1]]) for i in phrase_tokspans]
        return phrase_tokspans, phrases_vp

    """抽取名动词性短语"""

    def get_vnps(self, s):
        wds, postags = self.cut_wds(s)
        if not postags:
            return [], []
        if not (postags[-1].endswith("n") or postags[-1].endswith("l") or postags[-1].endswith("i")):
            return [], []
        phrase_tokspans = self.extract_sentgram(postags, self.VNP)
        if not phrase_tokspans:
            return [], []
        phrases_vnp = [''.join(wds[i[0]:i[1]]) for i in phrase_tokspans]
        phrase_tokspans2 = self.extract_sentgram(postags, self.NP)
        if not phrase_tokspans2:
            return [], []
        phrases_np = [''.join(wds[i[0]:i[1]]) for i in phrase_tokspans2]
        return phrases_vnp, phrases_np

    """提取短语"""

    def phrase_ip(self, content):
        import sys
        #print("原句：\n",content)
        spos = []
        events = []
        # 移除噪声数据
        content = self.remove_punc(content)
        # 分段落
        paras = self.split_paras(content)
        # 遍历循环每个段落
        for para in paras:
            # 将每个段落分割成若干个句子
            long_sents = self.split_long_sents(para)
            # 遍历循环每个句子
            for long_sent in long_sents:
                persons = []
                # 将每个句子进行分割成若干个短句
                short_sents = self.split_short_sents(long_sent)
                # 遍历循环每个短句
                for sent in short_sents:
                    # 分词：words为词；postags为词性
                    words, postags = self.cut_wds(sent)
                    # 人名实体识别
                    person = self.detect_person(words, postags)
                    # 处理代词等指示词语
                    words, postags = self.cite_resolution(words, postags, persons)
                    # 移除停用词
                    words, postags = self.clean_wds(words, postags)
                    ips = self.get_ips(words, postags)
                    #print(ips)
                    persons += person
                    # ips [(n,'分词'),(n,'分词')]
                    for ip in ips:
                    #   ip[1]取(n,'分词')中的分词
                        events.append(ip[1])
                        wds_tmp = []
                        postags_tmp = []
                        # 第二次分词，去除掉噪声、消岐、代词替换后
                        words, postags = self.cut_wds(ip[1])
                        # 识别动词短语
                        verb_tokspans, verbs = self.get_vps(words, postags)
                        #print("识别动词短语：",verb_tokspans,verbs)
                        # 识别出介宾短语
                        pp_tokspans, pps = self.get_pps(words, postags)
                        #print("识别介宾短语：",pp_tokspans,pps)
                        tmp_dict = {str(verb[0]) + str(verb[1]): ['V', verbs[idx]] for idx, verb in
                                    enumerate(verb_tokspans)}
                        #print(tmp_dict)
                        pp_dict = {str(pp[0]) + str(pp[1]): ['N', pps[idx]] for idx, pp in enumerate(pp_tokspans)}
                        #print(pp_dict)
                        tmp_dict.update(pp_dict)
                        #print(tmp_dict)
                        # sys.exit()
                        sort_keys = sorted([int(i) for i in tmp_dict.keys()])
                        for i in sort_keys:
                            if i < 10:
                                i = '0' + str(i)
                            wds_tmp.append(tmp_dict[str(i)][-1])
                            postags_tmp.append(tmp_dict[str(i)][0])
                        wds_tmp, postags_tmp = self.modify_duplicate(wds_tmp, postags_tmp, self.SPO_v, 'V')
                        wds_tmp, postags_tmp = self.modify_duplicate(wds_tmp, postags_tmp, self.SPO_n, 'N')
                        if len(postags_tmp) < 2:
                            continue
                        seg_index = []
                        i = 0
                        for wd, postag in zip(wds_tmp, postags_tmp):
                            if postag == 'V':
                                seg_index.append(i)
                            i += 1
                        spo = []
                        for indx, seg_indx in enumerate(seg_index):
                            if indx == 0:
                                pre_indx = 0
                            else:
                                pre_indx = seg_index[indx - 1]
                            if pre_indx < 0:
                                pre_indx = 0
                            if seg_indx == 0:
                                spo.append(('', wds_tmp[seg_indx], ''.join(wds_tmp[seg_indx + 1:])))
                            elif seg_indx > 0 and indx < 1:
                                spo.append(
                                    (''.join(wds_tmp[:seg_indx]), wds_tmp[seg_indx], ''.join(wds_tmp[seg_indx + 1:])))
                            else:
                                spo.append((''.join(wds_tmp[pre_indx + 1:seg_indx]), wds_tmp[seg_indx],
                                            ''.join(wds_tmp[seg_indx + 1:])))
                        spos += spo

        return events, spos


def del_all_graph(graph):
    #删除节点
    # 图中所有节点及关系都删除
    y = input("请确认是否要删除图库中所有节点及关系（y/n）：")
    if y == 'y':
        graph.delete_all()
        print("已确认删除图库所有节点和关系\n")
    elif y == 'n':
        print("已确认不删除图库所有节点和关系\n")
        pass
    else:
        print("=========请输入正确的提示引导=========\n")
        del_all_graph(graph)



if __name__ == '__main__':
    print("非结构化文本提取可视化程序开始")
    start_time = time.time()

    handler = ExtractEvent()

    # 连接图库
    graph = Graph('http://localhost:7474', auth=('neo4j', '123456789'))
    # 确认是否删除图库所有节点
    del_all_graph(graph)

    ''' 读取word文档 '''
    word = Document('D:\\Gitlab\\zixiao\\Knowledge_Graph\\Knowledge_Extract\\中共中央办公厅 国务院办公厅印发《关于加强金融服务民营企业的若干意见》.docx')
    # word = Document('D:\\Gitlab\\zixiao\\Knowledge_Graph\\Knowledge_Extract\\智器云.docx')

    # 得到文档内容
    word_paragraphs = word.paragraphs
    # 遍历文档列表，逐个进行处理
    for sen in word_paragraphs:
        # 输出此次处理语句
        print(sen.text)

        events, spos = handler.phrase_ip(sen.text)
        spos = [i for i in spos if i[0] and i[2]]
        # 输出提取的所有三元组关系
        print(spos)

        for word_list in spos:
            # 三元组输出
            print("准备写入图库当前三元组：",word_list)
            # 词性列表
            word_flag = []
            for i in range(len(word_list)):
                # 添加三元组中每项到词典
                jieba.add_word(word_list[i],freq='9999')
                # 对三元组中每个词进行分词
                word = jieba.posseg.cut(word_list[i])
                for i in word:
                    # 添加词性
                    word_flag.append(i.flag)


            # # 创建节点输出图库
            head = Node(f"{word_flag[0]}", name=f'{word_list[0]}')
            tail = Node(f"{word_flag[2]}", name=f'{word_list[2]}')

            # 匹配查找图库中节点
            matcher = NodeMatcher(graph)
            nodelist = list(matcher.match(f"{word_flag[0]}",name=f'{word_list[0]}'))
            if len(nodelist) > 0:
                # 表示节点存在，不需创建新的节点
                already_header = nodelist[0]  #
                # 可以直接添加关系
                entity = Relationship(already_header, f"{word_list[1]}", tail)
                graph.create(entity)
            else:
                # 表示图库中没有存在当前要写入的节点
                entity = Relationship(head, f"{word_list[1]}", tail)
                # 创建关系
                graph.create(entity)
            print("当前三元组写入结束\n")

        end_time = time.time()
        print(f"非结构化文本提取可视化程序结束，总耗时（秒）：{end_time-start_time}，写入图库节点数据量：{len(spos)}")




